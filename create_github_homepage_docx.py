#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("How to build global home page based on github setting v01.docx")
FONT_NAME = "Microsoft JhengHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(31, 77, 120)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def set_run_font(run, name=FONT_NAME, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"
    text_run = OxmlElement("w:r")
    text_run.append(text)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    paragraph._p.append(text_run)
    run = paragraph.add_run()
    run._r.append(end)


def add_page_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    label = p.add_run("第 ")
    set_run_font(label, size=9, color="5B677A")
    add_field(p, "PAGE")
    middle = p.add_run(" / ")
    set_run_font(middle, size=9, color="5B677A")
    add_field(p, "NUMPAGES")
    tail = p.add_run(" 頁")
    set_run_font(tail, size=9, color="5B677A")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.875), Inches(4.625)]
    hdr = table.rows[0].cells
    for idx, text in enumerate(["項目", "設定 / 意義"]):
        hdr[idx].width = widths[idx]
        set_cell_shading(hdr[idx], "E8EEF5")
        set_cell_text(hdr[idx], text, bold=True)
    rows = [
        ("GitHub 帳號", "greenmoon"),
        ("Repository", "greenmoon/codex_on_bed_project"),
        ("本機資料夾", "/Users/kevinwei/Dropbox/0DownLoad/codex_on_bed_project"),
        ("全球首頁 URL", "https://greenmoon.github.io/codex_on_bed_project/"),
        ("來源分支", "main"),
        ("發布資料夾", "/ (root)"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate([left, right]):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[idx], text, bold=(idx == 0))


def main():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("如何用 GitHub 設定建立全球首頁 v01")
    set_run_font(r, size=22, color="0B2545", bold=True)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    sr = sub.add_run("把本機 BED CSV viewer 發布成 GitHub Pages 全球網站的快速檢查文件。")
    set_run_font(sr, italic=True)
    add_code(doc, [
        "流程觀念：Python builder -> index.html -> GitHub repository -> GitHub Pages -> 全球瀏覽器網址",
    ])

    doc.add_heading("1. 完成結果", level=1)
    add_summary_table(doc)

    doc.add_heading("2. 本機專案流程", level=1)
    add_number(doc, "執行 Python builder，在本機建立或更新 viewer。")
    add_code(doc, [
        "cd /Users/kevinwei/Dropbox/0DownLoad/codex_on_bed_project",
        "python3 build_bed_curve_viewer.py",
    ])
    add_number(doc, "在瀏覽器檢查本機 viewer。")
    add_code(doc, [
        "python3 -m http.server 8765",
        "open http://localhost:8765/",
    ])
    add_number(doc, "Viewer 確認沒問題後，把產生的檔案 commit。")
    add_code(doc, [
        "git status",
        "git add .",
        "git commit -m \"Update BED curve viewer\"",
    ])

    doc.add_heading("3. GitHub repository 設定", level=1)
    add_bullet(doc, "建立 repository：greenmoon/codex_on_bed_project。")
    add_bullet(doc, "使用 HTTPS remote：https://github.com/greenmoon/codex_on_bed_project.git。")
    add_bullet(doc, "完成認證後，把本機 main branch push 上去。")
    add_code(doc, [
        "git remote set-url origin https://github.com/greenmoon/codex_on_bed_project.git",
        "git push -u origin main",
    ])

    doc.add_heading("4. Token 認證重點", level=1)
    add_bullet(doc, "GitHub 已不接受一般帳號密碼做 git push。")
    add_bullet(doc, "Terminal 問 Password 時，請貼上 Personal Access Token。")
    add_bullet(doc, "Fine-grained token 要選 repository：greenmoon/codex_on_bed_project。")
    add_bullet(doc, "Repository permissions 必須包含 Contents: Read and write。")
    add_bullet(doc, "如果 macOS 快取了錯誤 token，先清掉再重新 push。")
    add_code(doc, [
        "printf \"protocol=https\\nhost=github.com\\n\" | git credential-osxkeychain erase",
        "git push -u origin main",
    ])

    doc.add_heading("5. GitHub Pages 設定", level=1)
    add_bullet(doc, "開啟 repository 的 Settings，進入 Pages。")
    add_bullet(doc, "Source 設為 Deploy from a branch。")
    add_bullet(doc, "Branch 設為 main。")
    add_bullet(doc, "Folder 設為 / (root)。")
    add_bullet(doc, "按 Save，等待 GitHub Pages build 完成。")
    add_bullet(doc, "成功訊息會顯示：GitHub Pages source saved。")
    add_bullet(doc, "最後全球網址：https://greenmoon.github.io/codex_on_bed_project/。")

    doc.add_page_break()
    doc.add_heading("6. 常見錯誤與修正", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [Inches(2.3), Inches(4.2)]
    for idx, text in enumerate(["錯誤", "修正方式"]):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, text, bold=True)
    rows = [
        ("Password authentication is not supported", "使用 GitHub token，不要輸入 GitHub 帳號密碼。"),
        ("403 Permission denied", "重新產生 token，並確認選到此 repo，Contents 權限為 Read and write。"),
        ("Pages 顯示 404", "等待 1 到 3 分鐘再 refresh；也可到 Actions 看 deployment 狀態。"),
        ("本機 file list 沒更新", "重新執行 build_bed_curve_viewer.py，讓 index.html 嵌入最新資料夾檔案清單。"),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        for idx, text in enumerate([left, right]):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[idx], text)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
